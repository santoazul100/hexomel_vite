from pathlib import Path
from docx import Document
from pypdf import PdfReader
import json, re

docx_path = Path(r'C:/escola/pap/code/hexomel_vite/estudo/relatorio/analisei e acreentar/relatorio_v1.docx')
pdf_path = Path(r'C:/Users/rodrigo.silva/Downloads/A.17 RelatorioFinalPAP 1 - Artur Cruz (1).pdf')
out_dir = Path(r'C:/escola/pap/code/hexomel_vite/.codex_tmp/report_analysis')
out_dir.mkdir(parents=True, exist_ok=True)

result = {}
doc = Document(str(docx_path))
paras = []
for p in doc.paragraphs:
    text = ' '.join(p.text.split())
    if text:
        paras.append({'style': p.style.name if p.style else '', 'text': text})
tables = []
for t in doc.tables:
    rows=[]
    for row in t.rows:
        rows.append([' '.join(c.text.split()) for c in row.cells])
    tables.append(rows)
result['my_docx'] = {'paragraphs': paras, 'tables': tables}

reader = PdfReader(str(pdf_path))
pages=[]
for i,p in enumerate(reader.pages, start=1):
    try:
        txt = p.extract_text() or ''
    except Exception as e:
        txt = f'[EXTRACT_ERROR {e}]'
    # keep spaces compact
    txt = re.sub(r'[ \t]+', ' ', txt).strip()
    pages.append({'page': i, 'text': txt})
result['reference_pdf'] = {'pages': pages, 'page_count': len(pages)}

my_headings=[]
for x in paras:
    s=x['style'].lower(); text=x['text']
    if 'heading' in s or 'título' in s or 'title' in s or re.match(r'^(\d+(\.\d+)*)\s+', text):
        my_headings.append(text)
ref_text='\n'.join(p['text'] for p in pages)
ref_numbered = re.findall(r'(?<!\d)((?:[1-9]|1[0-9])(?:\.\d+)*\s+[A-ZÁÉÍÓÚÂÊÔÃÕÇ][^\n]{2,100})', ref_text)
seen=set(); ref_heads=[]
for h in ref_numbered:
    h=' '.join(h.split())
    h=re.sub(r'\s+\d+$','',h)
    if h not in seen:
        seen.add(h); ref_heads.append(h)

(out_dir/'my_docx_text.json').write_text(json.dumps(result['my_docx'], ensure_ascii=False, indent=2), encoding='utf-8')
(out_dir/'reference_pdf_text.json').write_text(json.dumps(result['reference_pdf'], ensure_ascii=False, indent=2), encoding='utf-8')
(out_dir/'summary.txt').write_text('MY_HEADINGS\n'+'\n'.join(my_headings[:250])+'\n\nREF_NUMBERED_HEADINGS\n'+'\n'.join(ref_heads[:300])+'\n\nPDF_PAGES '+str(len(pages)), encoding='utf-8')
print(out_dir/'summary.txt')
print('my paragraphs', len(paras), 'tables', len(tables), 'pdf pages', len(pages))
