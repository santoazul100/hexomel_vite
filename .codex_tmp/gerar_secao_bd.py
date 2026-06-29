from __future__ import annotations

import json
import math
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:/escola/pap/code/hexomel_vite")
OUT_DIR = ROOT / "estudo" / "relatorio" / "analisei e acreentar"
ASSET_DIR = OUT_DIR / "assets_bd"
ANALYSIS_PATH = ROOT / ".codex_tmp" / "db_full_analysis.json"
DOCX_PATH = OUT_DIR / "SECAO_3_6_EVOLUCAO_BASE_DADOS_HEXOMEL.docx"


def font(size: int, bold: bool = False):
    candidates = [
        r"C:/Windows/Fonts/calibrib.ttf" if bold else r"C:/Windows/Fonts/calibri.ttf",
        r"C:/Windows/Fonts/arialbd.ttf" if bold else r"C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def rounded(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def wrapped(draw, text, xy, max_width, fnt, fill=(30, 30, 30), line_gap=4):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        test = (current + " " + word).strip()
        if draw.textbbox((0, 0), test, font=fnt)[2] <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + line_gap
    return y


def make_initial_vs_final():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1800, 1050), "#FAFBFC")
    draw = ImageDraw.Draw(img)
    title_font = font(46, True)
    h_font = font(30, True)
    body_font = font(23)
    small_font = font(20)
    draw.text((70, 45), "Evolução da Base de Dados da Hexomel", font=title_font, fill="#1A4D2E")
    draw.text((70, 105), "Da loja online simples para uma plataforma completa", font=body_font, fill="#6B7280")

    left = (70, 170, 850, 930)
    right = (950, 170, 1730, 930)
    rounded(draw, left, 28, "#FFF7E6", "#E2A93B", 4)
    rounded(draw, right, 28, "#EEF8F1", "#1A4D2E", 4)
    draw.text((110, 215), "Modelo inicial pensado", font=h_font, fill="#7A4F00")
    draw.text((990, 215), "Modelo final implementado", font=h_font, fill="#1A4D2E")

    initial = [
        "Cliente",
        "Produto",
        "Categoria",
        "Encomenda",
        "Stock/Fatura",
    ]
    final_groups = [
        ("Comércio", "produto, categoria, origem, carrinho, item_carrinho, encomenda, item_encomenda, favoritos, avaliação"),
        ("Utilizadores", "cliente, password_recovery, upgrade_requests"),
        ("Apicultores", "workshop, reserva_workshop"),
        ("Comunidade", "pergunta, resposta, mensagem_privada, bloqueio, denúncia"),
        ("Aprender", "quiz_pergunta, quiz_score, aprender_facto, aprender_glossario"),
        ("Gestão", "interacao, site_slugs, site_settings, menu_nav, cms_content"),
    ]

    y = 300
    for item in initial:
        rounded(draw, (150, y, 770, y + 72), 18, "#FFFFFF", "#E2A93B", 2)
        draw.text((185, y + 20), item, font=body_font, fill="#3F2E00")
        y += 105
    draw.line((460, 375, 460, 780), fill="#C58A11", width=5)
    for node_y in [336, 441, 546, 651, 756]:
        draw.ellipse((448, node_y, 472, node_y + 24), fill="#E2A93B")

    y = 292
    for label, items in final_groups:
        rounded(draw, (1010, y, 1670, y + 82), 18, "#FFFFFF", "#9DCEAA", 2)
        draw.text((1038, y + 13), label, font=body_font, fill="#1A4D2E")
        wrapped(draw, items, (1185, y + 13), 445, small_font, "#374151", 2)
        y += 98

    draw.text((105, 865), "Resultado: modelo suficiente para uma loja básica.", font=body_font, fill="#7A4F00")
    draw.text((985, 865), "Resultado: modelo relacional para e-commerce, comunidade, aprendizagem e gestão.", font=small_font, fill="#1A4D2E")

    path = ASSET_DIR / "evolucao_bd_hexomel.png"
    img.save(path)
    return path


def make_module_map(analysis):
    img = Image.new("RGB", (2000, 1250), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_font = font(44, True)
    h_font = font(25, True)
    body_font = font(19)
    small_font = font(17)

    draw.text((80, 50), "Mapa modular da base de dados final", font=title_font, fill="#1A4D2E")
    draw.text((80, 105), f"{analysis['table_count']} tabelas analisadas | {analysis['fk_count']} relações | {analysis['route_count']} rotas no backend", font=body_font, fill="#6B7280")

    center = (1000, 620)
    rounded(draw, (790, 515, 1210, 725), 34, "#1A4D2E", "#0E2E1A", 3)
    draw.text((865, 565), "cliente", font=font(34, True), fill="#FFFFFF")
    draw.text((842, 615), "utilizador central", font=body_font, fill="#F8D56B")

    modules = [
        ("Comércio", "#FFF7E6", "#D69E2E", analysis["areas"]["Comércio"], (120, 220)),
        ("Segurança", "#EEF2FF", "#4F46E5", analysis["areas"]["Utilizadores e segurança"], (760, 220)),
        ("Apicultores", "#ECFDF5", "#059669", analysis["areas"]["Apicultores e workshops"], (1400, 220)),
        ("Comunidade", "#FEF2F2", "#DC2626", analysis["areas"]["Comunidade social"], (120, 780)),
        ("Aprender", "#F5F3FF", "#7C3AED", analysis["areas"]["Aprendizagem"], (760, 830)),
        ("Gestão e SEO", "#F3F4F6", "#374151", analysis["areas"]["Gestão e SEO"], (1400, 780)),
    ]
    for label, fill, border, names, (x, y) in modules:
        rounded(draw, (x, y, x + 500, y + 285), 26, fill, border, 4)
        draw.text((x + 30, y + 28), label, font=h_font, fill=border)
        draw.line((center[0], center[1], x + 250, y + 142), fill=border, width=4)
        ty = y + 78
        for name in names:
            rounded(draw, (x + 32, ty, x + 468, ty + 30), 8, "#FFFFFF", border, 1)
            draw.text((x + 48, ty + 5), name, font=small_font, fill="#111827")
            ty += 36

    draw.text((80, 1165), "Nota: o modelo final combina o SQL base com migrações automáticas presentes no backend/server.js.", font=body_font, fill="#6B7280")
    path = ASSET_DIR / "mapa_modular_bd_hexomel.png"
    img.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Secção 3.6 - Evolução da Base de Dados | Página ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, txt, end])


def add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_table(doc, rows, widths):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[c_idx])
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(26, 77, 46)
    doc.add_paragraph()


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.name = "Calibri"


def build_doc(analysis, img_evolution, img_map):
    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("3.6 Evolução da Base de Dados")
    r.font.name = "Calibri"
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = RGBColor(26, 77, 46)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(18)
    sr = sub.add_run("Secção pronta para inserir antes de “3.7 Ferramentas” no relatório Hexomel")
    sr.font.name = "Calibri"
    sr.font.size = Pt(11)
    sr.italic = True
    sr.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_heading("3.6 Evolução da Base de Dados", level=1)
    add_para(doc, "A base de dados foi uma das partes que mais evoluiu durante o desenvolvimento da Hexomel. Na fase inicial, o projeto podia ser representado por um modelo simples de loja online, centrado em clientes, produtos, categorias e encomendas. Essa primeira visão era útil para começar, mas tornou-se limitada quando o projeto passou a incluir perfis de apicultor, carrinho persistente, pagamentos, workshops, comunidade, chat privado, analytics e conteúdos educativos.")
    add_para(doc, "A análise ao código atual mostra que a Hexomel utiliza MySQL através do pacote mysql2/promise, com uma pool de ligações e um wrapper próprio com métodos get, all, run e exec. O ficheiro backend/hexomel_mysql.sql define a estrutura principal, enquanto o backend/server.js acrescenta migrações automáticas para funcionalidades que foram surgindo durante o desenvolvimento. Isto demonstra uma evolução prática: a base de dados não ficou parada no desenho inicial, mas foi sendo adaptada às necessidades reais da aplicação.")
    add_para(doc, f"No estado analisado, o modelo final contém {analysis['table_count']} tabelas identificadas entre o SQL base e as migrações do backend, com {analysis['fk_count']} relações por chave estrangeira e {analysis['route_count']} rotas Express associadas ao funcionamento geral do site. Estes números mostram que a base de dados deixou de suportar apenas uma loja e passou a sustentar uma plataforma completa.")

    doc.add_picture(str(img_evolution), width=Inches(6.5))
    caption(doc, "Figura 1. Comparação entre o modelo inicial pensado e o modelo final implementado.")

    doc.add_heading("3.6.1 Primeira Versão Pensada", level=2)
    add_para(doc, "A primeira versão da base de dados estava pensada para responder às necessidades mínimas de uma loja online. O objetivo era guardar clientes, produtos, categorias e encomendas, permitindo que um utilizador pudesse consultar produtos e realizar compras. Esta fase tinha uma lógica mais conceptual e serviu como ponto de partida para organizar as entidades principais.")
    add_bullet(doc, "Cliente: entidade responsável por guardar os dados básicos do utilizador.")
    add_bullet(doc, "Produto: entidade central do catálogo, com nome, preço, descrição e stock.")
    add_bullet(doc, "Categoria: forma simples de organizar produtos por tipo.")
    add_bullet(doc, "Encomenda: registo da compra realizada por um cliente.")
    add_bullet(doc, "Stock/Fatura: conceitos previstos inicialmente para controlo de quantidade e comprovativo.")
    add_para(doc, "Esta versão era suficiente para imaginar uma loja funcional, mas ainda não respondia a necessidades como autenticação segura, papéis de utilizador, favoritos, avaliações, carrinho persistente, workshops, moderação, comunidade ou dashboards.")

    doc.add_heading("3.6.2 Nova Versão Implementada", level=2)
    add_para(doc, "A versão final implementada é bastante mais completa. A tabela cliente passou a ser o ponto central do sistema, relacionando-se com encomendas, carrinho, favoritos, avaliações, pedidos de upgrade, workshops, perguntas, respostas, mensagens, denúncias, pontuações de quiz e interações. Esta centralidade permite representar diferentes papéis: cliente, apicultor e administrador.")
    add_para(doc, "A estrutura comercial foi reforçada com tabelas próprias para carrinho, item_carrinho, encomenda e item_encomenda. Esta separação é importante porque evita guardar vários conceitos numa só tabela e permite registar corretamente cada produto comprado, quantidade e preço unitário no momento da compra. Assim, se o preço de um produto mudar no futuro, o histórico da encomenda continua coerente.")
    add_para(doc, "A plataforma também ganhou áreas novas. Para os apicultores foram adicionadas tabelas de workshops e reservas. Para a comunidade foram criadas perguntas, respostas, mensagens privadas, bloqueios e denúncias. Para a secção Aprender foram adicionadas perguntas de quiz, pontuações, factos educativos e glossário. Para a gestão foram criadas tabelas de analytics, slugs, definições do site, menu dinâmico e conteúdos CMS.")

    doc.add_picture(str(img_map), width=Inches(6.5))
    caption(doc, "Figura 2. Mapa modular da base de dados final da Hexomel.")

    add_table(doc, [
        ["Módulo", "Tabelas principais", "Função"],
        ["Comércio", ", ".join(analysis["areas"]["Comércio"]), "Suporta catálogo, carrinho, encomendas, favoritos e avaliações."],
        ["Utilizadores e segurança", ", ".join(analysis["areas"]["Utilizadores e segurança"]), "Guarda perfis, recuperação de password e pedidos para se tornar apicultor."],
        ["Apicultores e workshops", ", ".join(analysis["areas"]["Apicultores e workshops"]), "Permite criar workshops e registar reservas dos utilizadores."],
        ["Comunidade social", ", ".join(analysis["areas"]["Comunidade social"]), "Gere perguntas, respostas, chat privado, bloqueios e denúncias."],
        ["Aprendizagem", ", ".join(analysis["areas"]["Aprendizagem"]), "Suporta factos educativos, glossário, quiz e leaderboard."],
        ["Gestão e SEO", ", ".join(analysis["areas"]["Gestão e SEO"]), "Apoia analytics, slugs, definições, menu e CMS."],
    ], [1700, 4300, 3360])

    doc.add_heading("3.6.3 Comparação e Melhorias Realizadas", level=2)
    add_para(doc, "A principal melhoria foi a passagem de um modelo básico para um modelo relacional modular. Em vez de uma base de dados limitada a produtos e compras, a Hexomel passou a ter uma estrutura capaz de acompanhar o funcionamento real do website. Isto permitiu separar responsabilidades, reduzir repetição de dados e representar melhor os fluxos da aplicação.")
    add_bullet(doc, "A relação entre encomendas e produtos foi normalizada através de item_encomenda.")
    add_bullet(doc, "O carrinho passou a ter tabelas próprias, permitindo persistência e sincronização.")
    add_bullet(doc, "Os produtos passaram a ter categoria, origem, apicultor, estado, tags, destaque e slug.")
    add_bullet(doc, "Os utilizadores passaram a ter papéis, perfil público, avatar, privacidade, verificação e 2FA de checkout.")
    add_bullet(doc, "A plataforma passou a suportar funcionalidades sociais, educativas e administrativas.")
    add_bullet(doc, "Foram adicionadas tabelas de suporte a SEO, CMS e analytics, aproximando o projeto de uma aplicação real.")

    add_table(doc, [
        ["Aspeto", "Modelo inicial", "Modelo final"],
        ["Objetivo", "Loja online simples.", "Plataforma completa de e-commerce, comunidade e aprendizagem."],
        ["Utilizadores", "Cliente básico.", "Cliente, apicultor e administrador com permissões diferentes."],
        ["Produtos", "Produto associado a categoria.", "Produto com categoria, origem, apicultor, tags, status, destaque e slug."],
        ["Compras", "Encomenda simples.", "Carrinho, itens, encomenda, itens de encomenda, envio, pagamento e recibos."],
        ["Comunidade", "Não prevista.", "Perguntas, respostas, chat, bloqueios e denúncias."],
        ["Educação", "Não prevista.", "Quiz, leaderboard, factos e glossário apícola."],
        ["Gestão", "Administração reduzida.", "CMS, menu dinâmico, slugs, analytics e moderação."],
    ], [1700, 3300, 4360])

    doc.add_heading("3.6.4 Análise crítica e melhorias recomendadas", level=2)
    add_para(doc, "A estrutura atual está adequada ao tamanho e ambição da PAP, mas a análise ao código permite identificar melhorias que poderiam ser aplicadas numa fase futura. A mais importante seria consolidar todas as migrações num sistema próprio de versionamento da base de dados, em vez de manter várias instruções CREATE TABLE e ALTER TABLE diretamente no server.js. Isto tornaria o backend mais limpo e permitiria saber exatamente que alterações foram aplicadas em cada versão.")
    add_bullet(doc, "Criar uma pasta de migrations numeradas, por exemplo 001_initial_schema.sql, 002_learning_tables.sql e 003_social_tables.sql.")
    add_bullet(doc, "Atualizar o ficheiro hexomel_mysql.sql para incluir também as tabelas criadas dinamicamente, como quiz_pergunta, quiz_score, aprender_facto, aprender_glossario, site_slugs, site_settings, menu_nav, password_recovery, cms_content e reserva_workshop.")
    add_bullet(doc, "Uniformizar nomes de colunas, evitando misturar português, inglês e pequenas gralhas como Data_Resgistro.")
    add_bullet(doc, "Adicionar índices em campos muito pesquisados, como Email, Slug, Status, ID_Apicultor e datas de encomenda/interação.")
    add_bullet(doc, "Separar melhor o backend por módulos ou routers, porque o ficheiro server.js concentra muitas responsabilidades.")
    add_para(doc, "Estas melhorias não retiram valor ao trabalho realizado. Pelo contrário, mostram que o projeto já atingiu uma dimensão em que práticas mais profissionais de organização seriam úteis. Para uma PAP, a evolução da base de dados demonstra aprendizagem, adaptação e capacidade de transformar uma ideia inicial num sistema funcional.")

    doc.add_heading("3.6.5 Síntese", level=2)
    add_para(doc, "A evolução da base de dados acompanha a evolução da própria Hexomel. A primeira versão respondia a uma loja simples; a versão final sustenta uma plataforma com vendas, perfis, apicultores, workshops, comunidade, chat, denúncias, aprendizagem, analytics e administração. Esta transformação é uma das partes mais importantes do projeto, porque mostra que a base de dados foi pensada em função das funcionalidades reais implementadas.")

    core = doc.core_properties
    core.title = "Secção 3.6 - Evolução da Base de Dados Hexomel"
    core.subject = "Relatório PAP Hexomel"
    core.author = "Rodrigo Filipe Costa Silva"
    core.keywords = "Hexomel, base de dados, MySQL, PAP, e-commerce"
    doc.save(DOCX_PATH)
    return DOCX_PATH


def main():
    analysis = json.loads(ANALYSIS_PATH.read_text(encoding="utf-8"))
    img_evolution = make_initial_vs_final()
    img_map = make_module_map(analysis)
    docx = build_doc(analysis, img_evolution, img_map)
    print(docx)
    print(img_evolution)
    print(img_map)


if __name__ == "__main__":
    main()
