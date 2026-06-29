from __future__ import annotations

import json
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:/escola/pap/code/hexomel_vite")
OUT_DIR = ROOT / "estudo" / "relatorio" / "analisei e acreentar"
ASSET_DIR = OUT_DIR / "assets_backend"
ANALYSIS_PATH = ROOT / ".codex_tmp" / "backend_45_analysis.json"
DOCX_PATH = OUT_DIR / "SECAO_4_5_BACKEND_HEXOMEL_MELHORADO.docx"

GREEN = "#1A4D2E"
GOLD = "#D69E2E"
LIGHT_GREEN = "#EEF8F1"
LIGHT_GOLD = "#FFF7E6"
BLUE = "#4F46E5"
RED = "#DC2626"
PURPLE = "#7C3AED"
GRAY = "#F3F4F6"
INK = "#111827"
MUTED = "#6B7280"


def font(size: int, bold: bool = False):
    candidates = [
        r"C:/Windows/Fonts/calibrib.ttf" if bold else r"C:/Windows/Fonts/calibri.ttf",
        r"C:/Windows/Fonts/arialbd.ttf" if bold else r"C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def rounded(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def wrap(draw, text, xy, max_width, fnt, fill=INK, gap=4):
    words = text.split()
    lines = []
    line = ""
    for word in words:
        test = (line + " " + word).strip()
        if draw.textbbox((0, 0), test, font=fnt)[2] <= max_width:
            line = test
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + gap
    return y


def arrow(draw, start, end, color="#374151", width=5):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=color, width=width)
    if x2 >= x1:
        draw.polygon([(x2, y2), (x2 - 18, y2 - 10), (x2 - 18, y2 + 10)], fill=color)
    else:
        draw.polygon([(x2, y2), (x2 + 18, y2 - 10), (x2 + 18, y2 + 10)], fill=color)


def make_backend_architecture(analysis):
    img = Image.new("RGB", (1900, 1200), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title = font(44, True)
    h = font(27, True)
    body = font(20)
    small = font(17)
    draw.text((70, 45), "Arquitetura do Backend Hexomel", font=title, fill=GREEN)
    draw.text((70, 100), "Organização funcional do servidor Node.js/Express e ligações externas", font=body, fill=MUTED)

    boxes = [
        ("Frontend", "Pedidos fetch enviados pelas páginas HTML e módulos JavaScript.", (80, 260, 395, 455), LIGHT_GOLD, GOLD),
        ("Express server.js", f"Servidor principal com {analysis['server_lines']} linhas e {analysis['routes_total']} rotas REST.", (560, 220, 1010, 500), "#EEF2FF", BLUE),
        ("Middlewares", "CORS, compression, JSON parser, uploads Multer, authenticateToken e isAdmin.", (1165, 260, 1580, 455), LIGHT_GREEN, GREEN),
        ("MySQL", "Pool mysql2/promise, wrapper db.get/all/run/exec e migrations automáticas.", (250, 690, 675, 925), "#F5F3FF", PURPLE),
        ("Serviços externos", "Stripe, Google OAuth e SMTP/Nodemailer para pagamentos, login e emails.", (760, 690, 1190, 925), "#FEF2F2", RED),
        ("Ficheiros públicos", "Uploads, imagens de produto e documentos submetidos por utilizadores.", (1275, 690, 1700, 925), GRAY, "#374151"),
    ]
    for name, text, box, fill, border in boxes:
        rounded(draw, box, 26, fill, border, 4)
        draw.text((box[0] + 28, box[1] + 28), name, font=h, fill=border)
        wrap(draw, text, (box[0] + 28, box[1] + 78), box[2] - box[0] - 56, body, INK, 5)
    arrow(draw, (395, 355), (560, 355), GOLD)
    arrow(draw, (1010, 355), (1165, 355), BLUE)
    arrow(draw, (785, 500), (465, 690), PURPLE)
    arrow(draw, (860, 500), (975, 690), RED)
    arrow(draw, (950, 500), (1475, 690), "#374151")
    rounded(draw, (110, 1030, 1790, 1100), 18, "#FFFFFF", "#D1D5DB", 2)
    draw.text((140, 1054), "Ideia-chave: o backend atua como camada de controlo entre a interface, a base de dados e os serviços externos, validando permissões e mantendo a lógica do negócio fora do frontend.", font=small, fill=INK)
    path = ASSET_DIR / "backend_arquitetura_hexomel.png"
    img.save(path)
    return path


def make_request_lifecycle():
    img = Image.new("RGB", (1900, 1420), "#FAFBFC")
    draw = ImageDraw.Draw(img)
    title = font(54, True)
    h = font(30, True)
    body = font(26)
    small = font(24)
    draw.text((70, 45), "Ciclo de um pedido no backend", font=title, fill=GREEN)
    draw.text((70, 115), "Fluxo principal desde a acao do utilizador ate a resposta final da API.", font=body, fill=MUTED)
    steps = [
        ("1. Acao no site", "Login, carrinho,\nmensagem ou\nconteudo."),
        ("2. Fetch API", "Pedido HTTP\npara uma rota\n/api/..."),
        ("3. Middleware", "CORS, JSON,\nJWT, admin\nou upload."),
        ("4. Logica Express", "Validacao,\nregras e\nqueries SQL."),
        ("5. BD/servicos", "MySQL, Stripe,\nGoogle OAuth\nou SMTP."),
        ("6. Resposta", "JSON, erro,\nrecibo ou\natualizacao."),
    ]
    colors = [GOLD, BLUE, GREEN, PURPLE, RED, "#374151"]
    center = (950, 705)
    circle_radius = 178
    draw.ellipse(
        (center[0] - circle_radius, center[1] - circle_radius, center[0] + circle_radius, center[1] + circle_radius),
        fill=LIGHT_GREEN,
        outline=GREEN,
        width=8,
    )
    draw.text((center[0] - 92, center[1] - 62), "BACKEND", font=font(38, True), fill=GREEN)
    draw.text((center[0] - 122, center[1] - 10), "Express.js", font=font(32, True), fill=INK)
    draw.text((center[0] - 106, center[1] + 38), "valida e", font=font(26), fill=MUTED)
    draw.text((center[0] - 112, center[1] + 74), "processa", font=font(26), fill=MUTED)

    boxes = [
        (120, 250, 470, 470),
        (775, 210, 1125, 430),
        (1430, 250, 1780, 470),
        (1430, 780, 1780, 1000),
        (775, 910, 1125, 1130),
        (120, 780, 470, 1000),
    ]
    anchors = [
        ((470, 360), (770, 610)),
        ((950, 430), (950, 525)),
        ((1430, 360), (1130, 610)),
        ((1430, 890), (1130, 800)),
        ((950, 910), (950, 885)),
        ((470, 890), (770, 800)),
    ]
    for idx, ((label, text), box, color) in enumerate(zip(steps, boxes, colors)):
        rounded(draw, box, 28, "#FFFFFF", color, 6)
        draw.text((box[0] + 28, box[1] + 24), label, font=h, fill=color)
        for line_idx, line in enumerate(text.split("\n")):
            draw.text((box[0] + 30, box[1] + 78 + line_idx * 38), line, font=small, fill=INK)
        arrow(draw, anchors[idx][0], anchors[idx][1], color, 5)

    rounded(draw, (135, 1190, 1765, 1345), 24, LIGHT_GREEN, GREEN, 3)
    draw.text((175, 1220), "Exemplo aplicado ao checkout", font=h, fill=GREEN)
    wrap(
        draw,
        "No checkout, o backend valida a sessao, confirma o carrinho, cria a encomenda, comunica com Stripe quando necessario e devolve ao frontend a informacao para continuar o pagamento ou apresentar o recibo.",
        (175, 1270),
        1540,
        body,
        INK,
        5,
    )
    path = ASSET_DIR / "backend_ciclo_pedido.png"
    img.save(path)
    return path

def make_routes_chart(analysis):
    img = Image.new("RGB", (1900, 1250), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title = font(44, True)
    h = font(25, True)
    body = font(19)
    big = font(38, True)
    draw.text((70, 45), "Distribuição das rotas do backend", font=title, fill=GREEN)
    draw.text((70, 100), "Leitura quantitativa das rotas Express identificadas em backend/server.js", font=body, fill=MUTED)
    stats = [
        ("GET", analysis["methods"].get("GET", 0), GREEN),
        ("POST", analysis["methods"].get("POST", 0), GOLD),
        ("PUT", analysis["methods"].get("PUT", 0), BLUE),
        ("PATCH", analysis["methods"].get("PATCH", 0), PURPLE),
        ("DELETE", analysis["methods"].get("DELETE", 0), RED),
    ]
    x = 100
    for label, value, color in stats:
        rounded(draw, (x, 190, x + 295, 335), 22, "#FFFFFF", color, 4)
        draw.text((x + 30, 220), str(value), font=big, fill=color)
        draw.text((x + 30, 282), label, font=body, fill=INK)
        x += 340
    draw.text((100, 430), "Rotas por área funcional", font=h, fill=GREEN)
    area_items = sorted(analysis["areas"].items(), key=lambda item: item[1], reverse=True)
    max_value = max(v for _, v in area_items)
    y = 500
    colors = [GREEN, GOLD, BLUE, PURPLE, RED, "#374151", "#059669", "#B45309"]
    for idx, (label, value) in enumerate(area_items):
        color = colors[idx % len(colors)]
        draw.text((120, y + 8), label, font=body, fill=INK)
        bar_width = int(1000 * value / max_value)
        rounded(draw, (560, y, 560 + bar_width, y + 38), 10, color, None, 1)
        draw.text((580 + bar_width, y + 8), str(value), font=body, fill=INK)
        y += 65
    rounded(draw, (1250, 455, 1780, 1010), 22, "#F9FAFB", "#CBD5E1", 2)
    draw.text((1285, 490), "Operações SQL no backend", font=h, fill=GREEN)
    y = 550
    for op, value in analysis["sql_ops"].items():
        draw.text((1290, y), op, font=body, fill=INK)
        draw.text((1650, y), str(value), font=body, fill=GREEN)
        y += 55
    path = ASSET_DIR / "backend_rotas_distribuicao.png"
    img.save(path)
    return path


def make_security_map():
    img = Image.new("RGB", (1900, 1450), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title = font(54, True)
    subtitle = font(28)
    h = font(32, True)
    body = font(26)
    draw.text((70, 45), "Seguranca e controlo de permissoes", font=title, fill=GREEN)
    draw.text((70, 115), "Camadas principais usadas pelo backend para proteger dados e rotas privadas.", font=subtitle, fill=MUTED)
    layers = [
        ("1. Password", "bcryptjs cria hashes das passwords\nantes de guardar na base de dados.", GOLD),
        ("2. Sessao", "jsonwebtoken gera e valida tokens JWT\npara identificar utilizadores autenticados.", BLUE),
        ("3. Middleware", "authenticateToken confirma o Bearer token\ne adiciona os dados do utilizador ao req.user.", GREEN),
        ("4. Admin", "isAdmin bloqueia rotas administrativas\nquando o utilizador nao tem perfil admin.", PURPLE),
        ("5. Checkout", "OTP por email reforca a validacao\nantes de finalizar uma compra.", RED),
        ("6. Moderacao", "Bloqueios, denuncias e restricoes\najudam a controlar a comunidade.", "#374151"),
    ]
    boxes = [
        (110, 230, 880, 430),
        (1020, 230, 1790, 430),
        (110, 520, 880, 720),
        (1020, 520, 1790, 720),
        (110, 810, 880, 1010),
        (1020, 810, 1790, 1010),
    ]
    for (name, text, color), box in zip(layers, boxes):
        rounded(draw, box, 28, "#F9FAFB", color, 6)
        draw.text((box[0] + 36, box[1] + 28), name, font=h, fill=color)
        for line_idx, line in enumerate(text.split("\n")):
            draw.text((box[0] + 38, box[1] + 88 + line_idx * 42), line, font=body, fill=INK)

    rounded(draw, (180, 1145, 1720, 1320), 28, LIGHT_GREEN, GREEN, 4)
    draw.text((225, 1180), "Ideia principal", font=h, fill=GREEN)
    wrap(
        draw,
        "O backend nao confia apenas no frontend. Antes de executar acoes sensiveis, valida a sessao, o papel do utilizador, os dados recebidos e as permissoes necessarias.",
        (225, 1240),
        1440,
        body,
        INK,
        8,
    )
    path = ASSET_DIR / "backend_seguranca_permissoes.png"
    img.save(path)
    return path

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("4.5 Backend Hexomel | Página ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, txt, end])


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.name = "Calibri"


def add_table(doc, rows, widths):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[c_idx])
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(26, 77, 46)
    doc.add_paragraph()


def build_doc(analysis, images):
    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("4.5 Desenvolvimento do Backend")
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = RGBColor(26, 77, 46)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(16)
    sr = sub.add_run("Secção melhorada para substituir o ponto 4.5 do relatório Hexomel")
    sr.font.name = "Calibri"
    sr.font.size = Pt(11)
    sr.italic = True
    sr.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_heading("4.5 Desenvolvimento do Backend", level=1)
    para(doc, "O backend da Hexomel é a camada responsável por ligar a interface do utilizador à base de dados e aos serviços externos. Enquanto o frontend apresenta páginas, formulários, animações e interações, o backend valida os dados recebidos, aplica regras de negócio, protege rotas privadas, comunica com o MySQL e devolve respostas em formato JSON para que a interface seja atualizada.")
    para(doc, f"No estado atual do projeto, o backend está concentrado principalmente no ficheiro backend/server.js, que possui cerca de {analysis['server_lines']} linhas e {analysis['routes_total']} rotas Express. Apesar de esta concentração tornar o ficheiro grande, ela mostra também a quantidade de funcionalidades implementadas: autenticação, catálogo, carrinho, checkout, encomendas, apicultores, workshops, administração, comunidade, chat, quiz, glossário, analytics e CMS.")

    doc.add_picture(str(images["architecture"]), width=Inches(6.5))
    caption(doc, "Figura 1. Arquitetura funcional do backend Hexomel.")

    doc.add_heading("4.5.1 Tecnologias utilizadas no backend", level=2)
    para(doc, "O backend foi desenvolvido em Node.js com Express.js. Esta escolha permitiu usar JavaScript em todo o projeto, tanto no cliente como no servidor, reduzindo a distância entre a lógica do frontend e a lógica da API. O Express foi usado para definir rotas HTTP e middlewares, enquanto o MySQL foi usado como sistema de gestão de base de dados relacional.")
    add_table(doc, [
        ["Tecnologia", "Função no backend"],
        ["Express.js", "Criação da API REST, definição de rotas e tratamento de pedidos HTTP."],
        ["mysql2/promise", "Ligação ao MySQL através de pool de conexões e queries assíncronas."],
        ["bcryptjs", "Hashing de passwords antes de guardar na base de dados."],
        ["jsonwebtoken", "Criação e validação de tokens JWT para sessões autenticadas."],
        ["Nodemailer", "Envio de emails de verificação, recuperação, OTP e recibos."],
        ["Stripe", "Criação e validação de sessões de pagamento."],
        ["google-auth-library", "Validação de login social com Google OAuth."],
        ["Multer", "Receção de uploads de imagens e documentos."],
        ["compression e cors", "Otimização de respostas e controlo de acesso entre frontend e backend."],
    ], [2200, 7160])

    doc.add_heading("4.5.2 Organização geral do servidor", level=2)
    para(doc, "O servidor é inicializado com Express e configurado para aceitar JSON, pedidos CORS, compressão e ficheiros estáticos. A ligação à base de dados é feita através de backend/config/db.js, onde é criada uma pool MySQL com limite de ligações. Esta abordagem é melhor do que abrir uma ligação nova para cada pedido, porque permite reutilizar conexões e responder a vários utilizadores de forma mais eficiente.")
    bullet(doc, "server.js: concentra a API REST, migrações, integrações e lógica principal do projeto.")
    bullet(doc, "config/env.js: lê variáveis de ambiente e centraliza configurações de base de dados.")
    bullet(doc, "config/db.js: cria a pool MySQL e fornece métodos db.get, db.all, db.run e db.exec.")
    bullet(doc, "middleware/auth.js: protege rotas com JWT e valida permissões de administrador.")

    doc.add_heading("4.5.3 API REST e distribuição de rotas", level=2)
    para(doc, "A API REST é a parte mais visível do backend. Cada rota representa uma funcionalidade que pode ser chamada pelo frontend. Por exemplo, quando o utilizador faz login, adiciona um produto ao carrinho, cria uma pergunta na comunidade ou altera uma encomenda no painel admin, o frontend envia um pedido para uma rota específica do backend.")
    doc.add_picture(str(images["routes"]), width=Inches(6.5))
    caption(doc, "Figura 2. Distribuição das rotas Express por método e área funcional.")
    add_table(doc, [
        ["Área", "N.º de rotas", "Responsabilidade"],
        ["Autenticação e conta", analysis["areas"]["Autenticação e conta"], "Login, registo, Google OAuth, perfil, recuperação e validação de sessão."],
        ["Catálogo e produtos", analysis["areas"]["Catálogo e produtos"], "Produtos, categorias, origens, slugs, avaliações e pesquisa."],
        ["Carrinho, checkout e encomendas", analysis["areas"]["Carrinho, checkout e encomendas"], "Carrinho, pagamento, encomendas, recibos e estado de checkout."],
        ["Apicultor e workshops", analysis["areas"]["Apicultor e workshops"], "Produtos do apicultor, bio, workshops, reservas e estatísticas."],
        ["Administração e CMS", analysis["areas"]["Administração e CMS"], "Gestão de utilizadores, produtos, encomendas, menu, CMS e moderação."],
        ["Comunidade e chat", analysis["areas"]["Comunidade e chat"], "Perguntas, respostas, mensagens privadas, bloqueios e denúncias."],
        ["Aprender e quiz", analysis["areas"]["Aprender e quiz"], "Factos educativos, glossário, perguntas de quiz, scores e leaderboard."],
    ], [2600, 1350, 5410])

    doc.add_heading("4.5.4 Ciclo de funcionamento de um pedido", level=2)
    para(doc, "O funcionamento de uma rota segue normalmente o mesmo ciclo. Primeiro o utilizador executa uma ação no site. Depois o frontend envia um pedido HTTP para a API. O backend aplica middlewares, valida dados, executa a lógica necessária, comunica com a base de dados ou com serviços externos e devolve uma resposta JSON ao frontend.")
    doc.add_picture(str(images["lifecycle"]), width=Inches(6.5))
    caption(doc, "Figura 3. Ciclo geral de um pedido no backend.")
    para(doc, "Este padrão é importante porque evita que o frontend tenha responsabilidade direta sobre dados sensíveis. O frontend apenas pede ações; o backend decide se elas são válidas, se o utilizador tem permissão e que alterações devem ser guardadas.")

    doc.add_heading("4.5.5 Ligação à base de dados", level=2)
    para(doc, "A ligação à base de dados é feita através de mysql2/promise, usando uma pool de conexões. O ficheiro db.js fornece quatro métodos principais: get para obter um registo, all para obter vários registos, run para executar alterações e exec para executar scripts SQL maiores. Esta abstração torna o restante código mais simples, porque as rotas não precisam de lidar diretamente com a criação da ligação.")
    bullet(doc, "db.get: usado quando se espera apenas um resultado, como procurar um utilizador pelo email.")
    bullet(doc, "db.all: usado para listas, como produtos, encomendas, mensagens ou perguntas.")
    bullet(doc, "db.run: usado para INSERT, UPDATE e DELETE.")
    bullet(doc, "db.exec: usado em scripts e migrações com várias instruções SQL.")
    para(doc, f"A análise ao server.js identificou muitas operações SQL: {analysis['sql_ops']['SELECT']} ocorrências de SELECT, {analysis['sql_ops']['INSERT']} de INSERT, {analysis['sql_ops']['UPDATE']} de UPDATE e {analysis['sql_ops']['DELETE']} de DELETE. Isto mostra que o backend está fortemente ligado à persistência de dados e que a aplicação depende de informação real guardada no MySQL.")

    doc.add_heading("4.5.6 Autenticação, permissões e segurança", level=2)
    para(doc, "A segurança do backend assenta principalmente em três elementos: passwords protegidas com bcryptjs, tokens JWT para sessões autenticadas e middlewares para proteger rotas privadas. Quando um utilizador faz login, o backend verifica a password e devolve um token. Nas rotas protegidas, esse token é enviado no cabeçalho Authorization e validado pelo middleware authenticateToken.")
    doc.add_picture(str(images["security"]), width=Inches(6.5))
    caption(doc, "Figura 4. Camadas de segurança e permissões no backend.")
    para(doc, "Além da autenticação normal, existe a função isAdmin, que impede utilizadores comuns de acederem a rotas administrativas. Esta separação é essencial, porque o projeto contém ações sensíveis como apagar produtos, gerir utilizadores, resolver denúncias, editar conteúdos CMS e consultar dashboards.")

    doc.add_heading("4.5.7 Integração com pagamentos, emails e login social", level=2)
    para(doc, "Uma parte importante do backend é a integração com serviços externos. O Stripe é usado para pagamentos, o Nodemailer para emails e o Google OAuth para autenticação social. Estas integrações tornam a Hexomel mais próxima de uma aplicação profissional, porque representam necessidades reais de uma plataforma de e-commerce.")
    bullet(doc, "Stripe: criação de sessões de checkout e verificação do estado do pagamento.")
    bullet(doc, "Nodemailer: envio de verificação de email, recuperação de password, OTP de checkout e recibos.")
    bullet(doc, "Google OAuth: login mais rápido e seguro através de conta Google.")
    bullet(doc, "Multer: upload de imagens de produtos e documentos associados a pedidos de upgrade.")

    doc.add_heading("4.5.8 Migrations e evolução do backend", level=2)
    para(doc, "Durante o desenvolvimento, o backend passou a incluir migrações automáticas no arranque do servidor. Estas migrações criam tabelas ou adicionam colunas quando novas funcionalidades são implementadas. Por exemplo, foram adicionadas estruturas para workshops, reservas, comunidade, mensagens privadas, denúncias, quiz, glossário, CMS, slugs e recuperação de password.")
    para(doc, "Esta abordagem ajudou durante a construção do projeto, porque permitiu evoluir rapidamente a base de dados. No entanto, numa versão futura mais profissional, estas migrações deveriam estar separadas em ficheiros próprios, numerados e controlados por versão.")

    doc.add_heading("4.5.9 Pontos fortes do backend", level=2)
    bullet(doc, "Integra várias áreas reais de uma aplicação: autenticação, e-commerce, admin, comunidade, chat e aprendizagem.")
    bullet(doc, "Usa base de dados relacional com pool MySQL e queries parametrizadas pelo mysql2.")
    bullet(doc, "Implementa autenticação JWT e validação de administrador.")
    bullet(doc, "Integra serviços externos relevantes: Stripe, Google OAuth e Nodemailer.")
    bullet(doc, "Gera emails/recibos e suporta uploads de ficheiros.")
    bullet(doc, "Tem rotas suficientes para demonstrar uma aplicação funcional e não apenas uma maquete.")

    doc.add_heading("4.5.10 Melhorias recomendadas", level=2)
    para(doc, "Apesar de funcional, o backend pode ser melhorado em termos de organização. O principal problema é a concentração de demasiadas responsabilidades no server.js. Como o ficheiro cresceu muito, torna-se mais difícil manter, encontrar erros e reutilizar código. Para uma versão futura, o ideal seria separar a API por módulos.")
    add_table(doc, [
        ["Problema atual", "Melhoria recomendada", "Benefício"],
        ["server.js demasiado grande", "Separar rotas em routers: auth.routes.js, products.routes.js, checkout.routes.js, admin.routes.js", "Código mais limpo e fácil de manter."],
        ["Migrações dentro do servidor", "Criar pasta migrations com ficheiros SQL numerados", "Evolução da BD mais controlada."],
        ["Regras misturadas com rotas", "Criar services para checkout, email, produtos e comunidade", "Reduz repetição e melhora testes."],
        ["Validações espalhadas", "Centralizar validações de inputs", "Menos erros e maior segurança."],
        ["Tratamento de erros repetido", "Criar middleware global de erros", "Respostas mais consistentes."],
    ], [2500, 3900, 2960])

    doc.add_heading("4.5.11 Síntese", level=2)
    para(doc, "O backend da Hexomel é uma das partes mais importantes do projeto, porque sustenta praticamente todas as funcionalidades dinâmicas do site. Ele gere autenticação, base de dados, compras, pagamentos, emails, administração, comunidade, chat e aprendizagem. Embora ainda possa ser melhor organizado, já demonstra competências avançadas de desenvolvimento web full-stack e aproxima a PAP de uma aplicação real.")

    core = doc.core_properties
    core.title = "Secção 4.5 - Desenvolvimento do Backend Hexomel"
    core.subject = "Relatório PAP Hexomel"
    core.author = "Rodrigo Filipe Costa Silva"
    core.keywords = "Hexomel, backend, Node.js, Express, MySQL, JWT, Stripe, PAP"
    doc.save(DOCX_PATH)
    return DOCX_PATH


def main():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    analysis = json.loads(ANALYSIS_PATH.read_text(encoding="utf-8"))
    images = {
        "architecture": make_backend_architecture(analysis),
        "lifecycle": make_request_lifecycle(),
        "routes": make_routes_chart(analysis),
        "security": make_security_map(),
    }
    docx = build_doc(analysis, images)
    print(docx)
    for path in images.values():
        print(path)


if __name__ == "__main__":
    main()
