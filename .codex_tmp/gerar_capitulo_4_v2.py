from __future__ import annotations

import json
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:/escola/pap/code/hexomel_vite")
OUT_DIR = ROOT / "estudo" / "relatorio" / "analisei e acreentar"
ASSET_DIR = OUT_DIR / "assets_cap4"
METRICS_PATH = ROOT / ".codex_tmp" / "cap4_metrics.json"
DOCX_PATH = OUT_DIR / "CAPITULO_4_DESENVOLVIMENTO_HEXOMEL_V2_COMPLEXO.docx"


GREEN = "#1A4D2E"
GOLD = "#D69E2E"
LIGHT_GOLD = "#FFF7E6"
LIGHT_GREEN = "#EEF8F1"
BLUE = "#2E74B5"
MUTED = "#6B7280"
INK = "#111827"
GRAY = "#F3F4F6"


def get_font(size: int, bold: bool = False):
    candidates = [
        r"C:/Windows/Fonts/calibrib.ttf" if bold else r"C:/Windows/Fonts/calibri.ttf",
        r"C:/Windows/Fonts/arialbd.ttf" if bold else r"C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def rounded(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_wrapped(draw, text, xy, max_width, font, fill=INK, line_gap=4):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = (current + " " + word).strip()
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + line_gap
    return y


def make_timeline(metrics):
    img = Image.new("RGB", (1900, 980), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title = get_font(44, True)
    h = get_font(25, True)
    body = get_font(19)
    small = get_font(17)
    draw.text((70, 45), "Evolução real do desenvolvimento Hexomel", font=title, fill=GREEN)
    draw.text((70, 100), "Leitura baseada no histórico Git: janeiro a junho de 2026", font=body, fill=MUTED)
    months = [
        ("Janeiro", "Base inicial", "Login, produtos, SGBD, perfil e primeiras áreas admin."),
        ("Fevereiro", "Catálogo + MySQL", "Logo, filtros, produto, conflitos resolvidos e ligação à BD."),
        ("Março", "Fluxo comercial", "Checkout multi-step, histórico de encomendas, apicultor e admin."),
        ("Abril", "Segurança", "2FA, email verification, dashboard admin e encomendas mais robustas."),
        ("Maio", "Experiência", "Workshops, curiosidades, 3D, educação e preparação da apresentação."),
        ("Junho", "Polimento final", "Aprender, leaderboard, CMS, diretório social, chat e correções UX."),
    ]
    x0, y0 = 120, 245
    gap = 285
    draw.line((170, 575, 170 + gap * 5, 575), fill=GOLD, width=8)
    for idx, (month, label, text) in enumerate(months):
        x = 100 + idx * gap
        draw.ellipse((x + 38, 540, x + 92, 594), fill=GOLD, outline="#8A5D00", width=3)
        rounded(draw, (x, y0 if idx % 2 == 0 else 625, x + 235, (y0 if idx % 2 == 0 else 625) + 245), 18, LIGHT_GOLD if idx % 2 == 0 else LIGHT_GREEN, GOLD if idx % 2 == 0 else GREEN, 3)
        y = y0 + 25 if idx % 2 == 0 else 650
        draw.text((x + 22, y), month, font=h, fill=GREEN)
        draw.text((x + 22, y + 38), label, font=small, fill="#7A4F00")
        draw_wrapped(draw, text, (x + 22, y + 75), 190, body, INK, 5)
    rounded(draw, (70, 875, 1830, 935), 16, GRAY, "#D1D5DB", 2)
    draw.text((100, 892), f"Métricas do projeto: {metrics['html_pages']} páginas HTML, {metrics['js_modules']} módulos JavaScript, {metrics['css_files']} ficheiros CSS, {metrics['routes_total']} rotas Express e {metrics['tables_total']} tabelas identificadas.", font=body, fill=INK)
    path = ASSET_DIR / "cap4_timeline_git.png"
    img.save(path)
    return path


def make_architecture(metrics):
    img = Image.new("RGB", (1800, 1200), "#FAFBFC")
    draw = ImageDraw.Draw(img)
    title = get_font(44, True)
    h = get_font(28, True)
    body = get_font(20)
    small = get_font(17)
    draw.text((70, 45), "Arquitetura funcional da Hexomel", font=title, fill=GREEN)
    draw.text((70, 100), "Separação por camadas: interface, API, lógica de negócio, dados e serviços externos", font=body, fill=MUTED)

    boxes = [
        ("Frontend", "HTML5 + CSS3 + JavaScript Vanilla\nVite, módulos ES6, UI, filtros, carrinho, perfil, admin, HexoHive", (90, 230, 500, 480), LIGHT_GOLD, GOLD),
        ("API REST", "Node.js + Express\n152 rotas, middlewares, validações, JSON, uploads e comunicação assíncrona", (700, 230, 1110, 480), "#EEF2FF", "#4F46E5"),
        ("Base de Dados", "MySQL + mysql2/promise\n28 tabelas, relações, migrações e persistência dos dados", (1310, 230, 1720, 480), LIGHT_GREEN, GREEN),
        ("Serviços Externos", "Stripe, Google OAuth, SMTP/Nodemailer, placeholders de imagem e Ngrok para testes", (395, 690, 845, 940), "#F5F3FF", "#7C3AED"),
        ("Administração", "CMS, dashboards, moderação, gestão de produtos, utilizadores, denúncias e conteúdos educativos", (955, 690, 1405, 940), "#FEF2F2", "#DC2626"),
    ]
    for name, desc, box, fill, border in boxes:
        rounded(draw, box, 26, fill, border, 4)
        draw.text((box[0] + 28, box[1] + 28), name, font=h, fill=border)
        draw_wrapped(draw, desc, (box[0] + 28, box[1] + 80), box[2] - box[0] - 56, body, INK, 6)
    arrows = [
        ((500, 355), (700, 355), "fetch/API"),
        ((1110, 355), (1310, 355), "SQL"),
        ((905, 480), (630, 690), "pagamentos/login/email"),
        ((905, 480), (1180, 690), "gestão interna"),
    ]
    for (x1, y1), (x2, y2), label in arrows:
        draw.line((x1, y1, x2, y2), fill="#374151", width=5)
        draw.polygon([(x2, y2), (x2 - 18 if x2 > x1 else x2 + 18, y2 - 10), (x2 - 18 if x2 > x1 else x2 + 18, y2 + 10)], fill="#374151")
        mx, my = (x1 + x2) // 2, (y1 + y2) // 2
        rounded(draw, (mx - 85, my - 18, mx + 85, my + 22), 10, "#FFFFFF", "#CBD5E1", 1)
        draw.text((mx - 70, my - 10), label, font=small, fill="#374151")
    rounded(draw, (120, 1035, 1680, 1110), 18, "#FFFFFF", "#D1D5DB", 2)
    draw.text((150, 1057), "Ponto forte: o projeto já não é só uma interface; tem lógica de negócio, autenticação, pagamentos, persistência, dashboards, comunidade e aprendizagem.", font=body, fill=INK)
    path = ASSET_DIR / "cap4_arquitetura_hexomel.png"
    img.save(path)
    return path


def make_user_flows():
    img = Image.new("RGB", (1900, 1250), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title = get_font(44, True)
    h = get_font(24, True)
    body = get_font(18)
    draw.text((70, 45), "Fluxos principais implementados", font=title, fill=GREEN)
    draw.text((70, 100), "Percursos de cliente, apicultor e administrador dentro da plataforma", font=body, fill=MUTED)
    flows = [
        ("Cliente", "#FFF7E6", "#D69E2E", ["Registo/Login", "Explora loja", "Adiciona ao carrinho", "2FA Checkout", "Stripe/MB Way", "Histórico/Recibo"]),
        ("Apicultor", "#EEF8F1", "#1A4D2E", ["Pedido upgrade", "Admin aprova", "Cria produtos", "Cria workshops", "Consulta stats", "Gere reservas"]),
        ("Administrador", "#EEF2FF", "#4F46E5", ["Dashboard", "Produtos/Utilizadores", "CMS/Aprender", "Encomendas", "Reports/Bloqueios", "Moderação final"]),
        ("Comunidade", "#FEF2F2", "#DC2626", ["Pergunta Q&A", "Resposta", "Votos", "Chat privado", "Bloqueio", "Denúncia"]),
    ]
    y = 205
    for title_txt, fill, border, steps in flows:
        rounded(draw, (80, y, 1820, y + 205), 24, fill, border, 3)
        draw.text((115, y + 28), title_txt, font=h, fill=border)
        x = 335
        for idx, step in enumerate(steps):
            rounded(draw, (x, y + 45, x + 205, y + 130), 18, "#FFFFFF", border, 2)
            draw_wrapped(draw, step, (x + 18, y + 70), 170, body, INK, 3)
            if idx < len(steps) - 1:
                draw.line((x + 205, y + 88, x + 245, y + 88), fill=border, width=4)
                draw.polygon([(x + 245, y + 88), (x + 230, y + 78), (x + 230, y + 98)], fill=border)
            x += 245
        y += 250
    path = ASSET_DIR / "cap4_fluxos_utilizador.png"
    img.save(path)
    return path


def make_module_dashboard(metrics):
    img = Image.new("RGB", (1900, 1180), "#FAFBFC")
    draw = ImageDraw.Draw(img)
    title = get_font(44, True)
    h = get_font(25, True)
    body = get_font(19)
    big = get_font(38, True)
    draw.text((70, 45), "Mapa de módulos e dimensão técnica", font=title, fill=GREEN)
    draw.text((70, 100), "Resumo quantitativo do projeto analisado no repositório", font=body, fill=MUTED)
    stats = [
        ("Páginas HTML", metrics["html_pages"], GOLD),
        ("Módulos JS", metrics["js_modules"], GREEN),
        ("Ficheiros CSS", metrics["css_files"], "#4F46E5"),
        ("Rotas API", metrics["routes_total"], "#DC2626"),
        ("Tabelas BD", metrics["tables_total"], "#7C3AED"),
        ("Migrações ALTER", metrics["alter_columns"], "#374151"),
    ]
    x, y = 90, 190
    for label, value, color in stats:
        rounded(draw, (x, y, x + 260, y + 145), 22, "#FFFFFF", color, 4)
        draw.text((x + 30, y + 28), str(value), font=big, fill=color)
        draw.text((x + 30, y + 90), label, font=body, fill=INK)
        x += 295
    areas = metrics["areas"]
    labels = [("Auth", areas["auth"]), ("Produtos", areas["products"]), ("Checkout", areas["checkout"]), ("Apicultor", areas["apicultor"]), ("Admin", areas["admin"]), ("Social", areas["social"]), ("Aprender", areas["learn"])]
    max_val = max(v for _, v in labels)
    draw.text((100, 420), "Distribuição de rotas por área", font=h, fill=GREEN)
    y = 480
    for label, value in labels:
        draw.text((120, y + 8), label, font=body, fill=INK)
        bar_w = int(1050 * value / max_val)
        rounded(draw, (300, y, 300 + bar_w, y + 38), 10, GOLD if label in {"Produtos", "Checkout"} else GREEN, None, 1)
        draw.text((320 + bar_w, y + 8), str(value), font=body, fill=INK)
        y += 65
    rounded(draw, (1200, 430, 1780, 935), 24, "#FFFFFF", "#D1D5DB", 2)
    draw.text((1235, 465), "Módulos frontend relevantes", font=h, fill=GREEN)
    important = ["api.js", "auth.js", "cart.js", "checkout.js", "shop.js", "produto.js", "admin.js", "dashboard-apicultor.js", "rede-social.js", "aprender.js", "analytics.js", "beeAnimation.js"]
    y = 520
    for idx, name in enumerate(important):
        col = 0 if idx < 6 else 1
        row = idx if idx < 6 else idx - 6
        xx = 1235 + col * 270
        yy = 520 + row * 58
        rounded(draw, (xx, yy, xx + 230, yy + 36), 10, "#F9FAFB", "#CBD5E1", 1)
        draw.text((xx + 12, yy + 9), name, font=body, fill=INK)
    path = ASSET_DIR / "cap4_modulos_metricas.png"
    img.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Capítulo 4 - Desenvolvimento do Projeto | Página ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, txt, end])


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.name = "Calibri"


def add_table(doc, rows, widths):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[c_idx])
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(26, 77, 46)
    doc.add_paragraph()
    return table


def build_doc(metrics, images):
    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("4. Desenvolvimento do Projeto")
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = RGBColor(26, 77, 46)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(16)
    sr = sub.add_run("Versão V2 expandida, com análise técnica, fluxos, tabelas e diagramas")
    sr.font.name = "Calibri"
    sr.font.size = Pt(11)
    sr.italic = True
    sr.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_heading("4. Desenvolvimento do Projeto", level=1)
    para(doc, "Este capítulo apresenta o desenvolvimento da Hexomel de forma mais aprofundada, combinando a evolução temporal do projeto, a arquitetura técnica, a organização do código, os fluxos de utilização e as principais decisões tomadas durante a implementação. O objetivo é demonstrar que a Hexomel não é apenas um conjunto de páginas estáticas, mas uma aplicação web completa, com frontend, backend, base de dados, autenticação, pagamentos, gestão administrativa, comunidade e componente educativa.")
    para(doc, f"A análise ao repositório mostra uma dimensão considerável para uma PAP: {metrics['html_pages']} páginas HTML, {metrics['js_modules']} módulos JavaScript, {metrics['css_files']} ficheiros CSS, {metrics['routes_total']} rotas Express e {metrics['tables_total']} tabelas identificadas entre o SQL base e as migrações do backend. Estes valores ajudam a perceber a complexidade real do projeto e justificam a necessidade de explicar o desenvolvimento por fases.")

    doc.add_picture(str(images["timeline"]), width=Inches(6.5))
    caption(doc, "Figura 1. Evolução real do desenvolvimento, organizada a partir do histórico Git.")

    doc.add_heading("4.1 Planeamento e calendarização real", level=2)
    para(doc, "Embora o planeamento inicial previsse uma loja online de mel relativamente simples, o desenvolvimento real tornou-se mais ambicioso. O histórico Git mostra uma evolução contínua entre janeiro e junho de 2026, começando pela base do site e terminando em funcionalidades avançadas como chat privado, leaderboard, CMS educativo, moderação social e correções de experiência de utilizador.")
    add_table(doc, [
        ["Período", "Foco técnico", "Resultado alcançado"],
        ["Janeiro", "Estrutura base", "Primeiras páginas, login, produtos, perfil, SGBD e área administrativa inicial."],
        ["Fevereiro", "Catálogo e base de dados", "Integração com MySQL, filtros, logo, páginas de produto e resolução de conflitos."],
        ["Março", "Fluxo comercial", "Checkout persistente, histórico de encomendas, perfil completo, apicultor e admin."],
        ["Abril", "Segurança e gestão", "2FA de checkout, verificação por email, dashboard admin e melhorias nas encomendas."],
        ["Maio", "Experiência e multimédia", "Workshops, curiosidades, elementos 3D, página educativa e preparação da apresentação."],
        ["Junho", "Polimento e funcionalidades finais", "Redesign Aprender, quiz, leaderboard, CMS, rede social, chat e correções finais."],
    ], [1450, 2700, 5210])

    doc.add_heading("4.2 Investigação e decisões técnicas", level=2)
    para(doc, "A pesquisa técnica foi feita em paralelo com o desenvolvimento. No início, o foco esteve em compreender a estrutura essencial de uma loja online: catálogo, utilizadores, carrinho, encomendas e base de dados. Depois, com a evolução do projeto, foi necessário estudar autenticação por JWT, hashing de passwords, envio de emails, pagamentos com Stripe, login social com Google OAuth, dashboards, moderação, uploads e persistência de dados.")
    bullet(doc, "A escolha de JavaScript Vanilla permitiu compreender o funcionamento real do DOM, eventos, módulos ES6 e Fetch API.")
    bullet(doc, "O Vite foi usado para tornar o desenvolvimento frontend mais rápido e organizado.")
    bullet(doc, "O backend em Node.js/Express centralizou a API REST e a lógica de negócio.")
    bullet(doc, "O MySQL foi escolhido por ser adequado a relações entre clientes, produtos, encomendas, workshops, mensagens e analytics.")
    bullet(doc, "Stripe, Nodemailer e Google OAuth aproximaram a PAP de um sistema real usado em produção.")

    doc.add_picture(str(images["architecture"]), width=Inches(6.5))
    caption(doc, "Figura 2. Arquitetura funcional da aplicação Hexomel.")

    doc.add_heading("4.3 Arquitetura geral do sistema", level=2)
    para(doc, "A Hexomel segue uma arquitetura em camadas. A camada de apresentação é composta por HTML, CSS e JavaScript no frontend; a camada de lógica de negócio está no servidor Node.js com Express; a camada de dados utiliza MySQL; e a camada de serviços externos integra Stripe, Google OAuth e SMTP/Nodemailer. Esta separação tornou o projeto mais organizado e facilitou a evolução das funcionalidades.")
    add_table(doc, [
        ["Camada", "Tecnologias/Ficheiros", "Responsabilidade"],
        ["Frontend", "HTML, CSS, JavaScript Vanilla, Vite", "Interface, navegação, formulários, filtros, carrinho, dashboards e interações."],
        ["Backend", "Node.js, Express, middlewares, server.js", "Rotas REST, autenticação, validação, pagamentos, emails, uploads e regras de negócio."],
        ["Base de dados", "MySQL, mysql2/promise, SQL e migrações", "Persistência de utilizadores, produtos, encomendas, comunidade, aprendizagem e analytics."],
        ["Serviços externos", "Stripe, Google OAuth, SMTP/Nodemailer", "Pagamentos, login social, emails transacionais e recibos."],
    ], [1600, 3100, 4660])

    doc.add_heading("4.4 Organização do frontend", level=2)
    para(doc, "O frontend foi organizado por páginas e módulos. Cada página HTML representa uma área do site, enquanto os ficheiros JavaScript em frontend/src implementam a lógica específica de cada funcionalidade. Esta divisão permitiu desenvolver componentes de forma mais isolada: a loja é tratada por shop.js, o carrinho por cart.js, o checkout por checkout.js, a administração por admin.js, o dashboard do apicultor por dashboard-apicultor.js e a rede social por rede-social.js.")
    para(doc, "A interface foi trabalhada com uma identidade visual ligada ao mel e à natureza, usando tons dourados, verdes e brancos. Também foram adicionados skeleton loaders, toasts, modais, animações e elementos 3D, o que reforça a vertente multimédia do curso.")

    doc.add_heading("4.5 Organização do backend e API", level=2)
    para(doc, f"O backend concentra {metrics['routes_total']} rotas Express, distribuídas por autenticação, produtos, checkout, apicultor, administração, comunidade e aprendizagem. Esta quantidade de rotas mostra que o projeto deixou de ser apenas visual: cada área do site comunica com o servidor e depende de dados reais.")
    add_table(doc, [
        ["Área da API", "N.º de rotas", "Exemplos de uso"],
        ["Autenticação", str(metrics["areas"]["auth"]), "Registo, login, Google OAuth, email verification e recuperação de password."],
        ["Produtos", str(metrics["areas"]["products"]), "Catálogo, categorias, origens, slugs, avaliações e produtos do apicultor."],
        ["Checkout", str(metrics["areas"]["checkout"]), "Carrinho, encomendas, pagamento, estado da sessão e recibos."],
        ["Apicultor", str(metrics["areas"]["apicultor"]), "Produtos, workshops, reservas, estatísticas e bio profissional."],
        ["Admin", str(metrics["areas"]["admin"]), "Utilizadores, produtos, encomendas, CMS, reports, blocos e moderação."],
        ["Social", str(metrics["areas"]["social"]), "Comunidade, membros, chat privado, bloqueios e denúncias."],
        ["Aprender", str(metrics["areas"]["learn"]), "Quiz, leaderboard, factos educativos e glossário."],
    ], [1800, 1400, 6160])

    doc.add_picture(str(images["modules"]), width=Inches(6.5))
    caption(doc, "Figura 3. Mapa de módulos e dimensão técnica do projeto.")

    doc.add_heading("4.6 Desenvolvimento da base de dados", level=2)
    para(doc, "A base de dados evoluiu em conjunto com as funcionalidades. Inicialmente, bastaria guardar clientes, produtos, categorias e encomendas. Porém, à medida que a Hexomel cresceu, foram necessárias tabelas para carrinho, itens de encomenda, favoritos, avaliações, workshops, reservas, comunidade, mensagens privadas, denúncias, bloqueios, analytics, quiz, glossário, factos educativos, CMS, slugs e configurações.")
    para(doc, "Esta evolução mostra uma passagem clara de uma loja simples para uma plataforma relacional mais completa. A utilização de chaves estrangeiras, tabelas intermédias e campos específicos para histórico de compra permitiu representar melhor os fluxos reais da aplicação.")

    doc.add_heading("4.7 Segurança e autenticação", level=2)
    para(doc, "A segurança foi implementada em várias camadas. As passwords são protegidas com bcryptjs, as sessões usam JWT, as rotas privadas passam por middleware de autenticação e as ações administrativas exigem validação de permissões. Para além disso, o checkout inclui verificação adicional por email, através de OTP temporário, aumentando a confiança no momento da compra.")
    bullet(doc, "Cliente: compra produtos, gere perfil, favoritos, encomendas e participa na comunidade.")
    bullet(doc, "Apicultor: gere produtos, workshops, reservas e estatísticas próprias.")
    bullet(doc, "Administrador: gere plataforma, conteúdos, utilizadores, encomendas, reports e moderação.")

    doc.add_picture(str(images["flows"]), width=Inches(6.5))
    caption(doc, "Figura 4. Fluxos principais de cliente, apicultor, administrador e comunidade.")

    doc.add_heading("4.8 Desenvolvimento do fluxo de compra", level=2)
    para(doc, "O fluxo de compra foi uma das partes mais importantes do projeto. O utilizador pesquisa produtos, filtra o catálogo, consulta a página de detalhe, adiciona itens ao carrinho, valida os dados de envio, confirma a sessão por 2FA e conclui o pagamento. O sistema evita criar encomendas desnecessárias demasiado cedo e procura sincronizar carrinho local e base de dados.")
    para(doc, "No backend, a encomenda é separada dos itens da encomenda. Esta decisão é importante porque permite guardar cada produto comprado com quantidade e preço unitário da altura, evitando erros históricos caso o produto mude de preço mais tarde.")

    doc.add_heading("4.9 Desenvolvimento da área do apicultor", level=2)
    para(doc, "A área do apicultor foi criada para que o produtor pudesse ter presença ativa na plataforma. O apicultor pode gerir os seus produtos, criar workshops e consultar estatísticas. Esta funcionalidade torna a Hexomel mais próxima de um marketplace especializado do que de uma loja controlada apenas por um administrador.")

    doc.add_heading("4.10 Desenvolvimento da área administrativa", level=2)
    para(doc, "A administração funciona como um CMS personalizado. O administrador consegue gerir produtos, categorias, origens, utilizadores, encomendas, pedidos de upgrade, conteúdos das páginas, menu, slugs, quiz, glossário, factos educativos e reports. Esta área é essencial porque permite manter o site dinâmico sem alterar diretamente o código.")
    para(doc, "Os dashboards com Chart.js ajudam a transformar dados em informação visual: receita, produtos mais vendidos, distribuição por categoria, estados de encomendas, crescimento de utilizadores e performance de parceiros. Esta camada analítica aproxima o projeto de sistemas reais de gestão.")

    doc.add_heading("4.11 Desenvolvimento da comunidade HexoHive", level=2)
    para(doc, "A componente social da Hexomel foi uma das funcionalidades mais avançadas. A comunidade inclui perguntas e respostas, diretório público de membros, perfis públicos, chat privado, bloqueios e denúncias. Esta área aumenta a interação entre clientes, apicultores e administradores, criando um ecossistema à volta do tema da apicultura.")
    para(doc, "O chat privado usa polling automático para atualizar conversas, marca mensagens como lidas e verifica bloqueios antes de permitir envio. A existência de denúncias e restrições de escrita mostra preocupação com moderação e segurança da comunidade.")

    doc.add_heading("4.12 Desenvolvimento da secção Aprender", level=2)
    para(doc, "A secção Aprender dá ao projeto uma componente educativa. Foram implementados factos interativos, glossário apícola, quiz gamificado e leaderboard. Esta área torna o site mais rico porque não se limita a vender produtos: também ensina conceitos relacionados com mel, abelhas e apicultura.")

    doc.add_heading("4.13 Testes, correções e iteração", level=2)
    para(doc, "Os testes foram feitos de forma progressiva. Sempre que uma funcionalidade era adicionada, eram verificados fluxos como login, carrinho, checkout, perfil, administração, criação de workshops, perguntas da comunidade, envio de mensagens e acesso a rotas protegidas. O histórico Git mostra várias fases de correção, incluindo problemas de layout, autenticação, encomendas e experiência visual.")
    add_table(doc, [
        ["Tipo de teste", "Exemplo aplicado", "Objetivo"],
        ["Funcional", "Adicionar produto ao carrinho e concluir checkout", "Garantir que o fluxo comercial funciona."],
        ["Segurança", "Aceder a rotas admin sem permissões", "Confirmar proteção por JWT e role."],
        ["Interface", "Testar header, dashboard e rede social", "Evitar sobreposições, quebras e inconsistências visuais."],
        ["Dados", "Criar encomenda, item_encomenda e recibo", "Confirmar persistência correta no MySQL."],
        ["Comunidade", "Enviar mensagem, bloquear e denunciar", "Validar moderação e segurança social."],
    ], [1800, 3600, 3960])

    doc.add_heading("4.14 Dificuldades encontradas", level=2)
    para(doc, "As principais dificuldades estiveram relacionadas com a integração entre várias áreas do sistema. Um erro no backend podia afetar o frontend; uma alteração na base de dados podia exigir atualização nas rotas; e uma nova funcionalidade de interface podia precisar de dados adicionais. A integração com Stripe, email, Google OAuth, permissões e MySQL exigiu testes constantes.")
    bullet(doc, "Manter o server.js organizado tornou-se difícil porque o projeto ganhou muitas rotas e responsabilidades.")
    bullet(doc, "A base de dados precisou de várias migrações à medida que surgiam novas funcionalidades.")
    bullet(doc, "A experiência visual exigiu muitas correções de layout, principalmente em páginas mais complexas.")
    bullet(doc, "A comunidade e o chat exigiram atenção extra a bloqueios, denúncias, mensagens lidas e permissões.")

    doc.add_heading("4.15 Aprendizagens técnicas", level=2)
    para(doc, "O desenvolvimento da Hexomel permitiu consolidar competências de frontend, backend, base de dados, integração de APIs, UX/UI, segurança e organização de projeto. A maior aprendizagem foi perceber que uma aplicação real depende da ligação entre várias partes: não basta fazer uma página bonita; é necessário garantir que os dados circulam corretamente, que o utilizador recebe feedback, que as permissões são respeitadas e que os estados são guardados.")
    bullet(doc, "Frontend modular com JavaScript Vanilla e Fetch API.")
    bullet(doc, "Backend REST com Node.js, Express e middlewares.")
    bullet(doc, "Modelo relacional MySQL ajustado aos fluxos reais da aplicação.")
    bullet(doc, "Integração com serviços externos como Stripe, Google OAuth e Nodemailer.")
    bullet(doc, "Uso de Git para acompanhar evolução, regressões e fases de desenvolvimento.")

    doc.add_heading("4.16 Síntese do capítulo", level=2)
    para(doc, "O desenvolvimento da Hexomel foi progressivo e iterativo. A versão final demonstra uma evolução significativa desde a ideia inicial de loja online até uma plataforma completa com e-commerce, comunidade, aprendizagem, administração, dashboards, segurança e pagamentos. Esta complexidade torna o projeto adequado a uma Prova de Aptidão Profissional, porque demonstra competências técnicas e capacidade de resolver problemas reais de desenvolvimento web.")

    core = doc.core_properties
    core.title = "Capítulo 4 - Desenvolvimento do Projeto Hexomel V2"
    core.subject = "Relatório PAP Hexomel"
    core.author = "Rodrigo Filipe Costa Silva"
    core.keywords = "Hexomel, PAP, desenvolvimento, Node.js, MySQL, frontend, backend"
    doc.save(DOCX_PATH)
    return DOCX_PATH


def main():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    metrics = json.loads(METRICS_PATH.read_text(encoding="utf-8"))
    images = {
        "timeline": make_timeline(metrics),
        "architecture": make_architecture(metrics),
        "flows": make_user_flows(),
        "modules": make_module_dashboard(metrics),
    }
    docx = build_doc(metrics, images)
    print(docx)
    for path in images.values():
        print(path)


if __name__ == "__main__":
    main()
